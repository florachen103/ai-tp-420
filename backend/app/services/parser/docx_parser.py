"""Word (.docx) 解析：按标题层级切分章节，保留表格文本。"""
from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn

from app.models.course import MaterialType
from app.services.parser.base import ParsedDocument, ParsedSection


def parse_docx(file_path: str, filename: str) -> ParsedDocument:
    doc = Document(file_path)
    sections: list[ParsedSection] = []
    current_title = "正文"
    current_buf: list[str] = []
    order = 0

    def flush():
        nonlocal order
        text = "\n".join(b for b in current_buf if b.strip())
        if text.strip():
            sections.append(ParsedSection(
                title=current_title,
                content=text,
                order_idx=order,
            ))
            order += 1

    for para in doc.paragraphs:
        style = (para.style.name or "").lower()
        text = para.text.strip()
        if not text:
            continue
        if style.startswith("heading") or style.startswith("标题"):
            flush()
            current_title = text
            current_buf = []
        else:
            current_buf.append(text)

    # 表格转成 markdown 形式追加到最近一节
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            current_buf.append("\n【表格】\n" + "\n".join(rows))

    flush()

    return ParsedDocument(
        filename=filename,
        material_type=MaterialType.WORD,
        sections=sections,
        meta={"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)},
    )
